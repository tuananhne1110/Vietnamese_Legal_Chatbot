from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import Dict, Any, Optional
import requests
import io
from docx import Document
import tempfile
import os

router = APIRouter(prefix="/api/ct01", tags=["CT01"])

class CT01FormData(BaseModel):
    formData: Dict[str, Any]
    cccdData: Optional[Dict[str, Any]] = None
    template: Optional[Dict[str, Any]] = None
    type: str = "docx"

class CT01SubmitData(BaseModel):
    formData: Dict[str, Any]
    cccdData: Optional[Dict[str, Any]] = None

@router.post("/generate")
async def generate_ct01_file(data: CT01FormData):
    """
    Tạo file CT01 từ template HTML với data đã điền
    """
    try:
        print(f"Received data: {data.formData}")
        print(f"Template info: {data.template}")
        
        # Đọc file ct01.html template
        template_path = "templates/ct01.html"
        print(f"Reading HTML template from: {template_path}")
        
        with open(template_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        print(f"HTML template loaded successfully: {len(html_content)} characters")
        
        # Điền data vào HTML template với cccd_data
        filled_html = fill_html_template_with_data(html_content, data.formData, data.cccdData)
        print(f"Data filled into HTML template")
        
        if data.type == "html":
            # Trả về HTML
            from fastapi.responses import Response
            filename = f"CT01-{data.template.get('code', 'CT01')}.html"
            return Response(
                content=filled_html,
                media_type="text/html",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        else:
            # Convert HTML to PDF hoặc DOCX
            output_file = convert_html_to_format(filled_html, data.type)
            
            # Trả về file
            from fastapi.responses import StreamingResponse
            
            filename = f"CT01-{data.template.get('code', 'CT01')}.{data.type}"
            media_type = "application/pdf" if data.type == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
            print(f"Returning file: {filename}")
            
            return StreamingResponse(
                io.BytesIO(output_file),
                media_type=media_type,
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        
    except Exception as e:
        print(f"Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating file: {str(e)}")

@router.post("/submit")
async def submit_ct01_form(data: CT01SubmitData):
    """
    Nộp form CT01 trực tuyến
    """
    try:
        # TODO: Implement online submission logic
        # Có thể lưu vào database, gửi email, etc.
        
        return {
            "success": True,
            "message": "Form submitted successfully",
            "reference_id": f"CT01-{len(data.formData)}-{hash(str(data.formData)) % 10000}"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error submitting form: {str(e)}")

def fill_html_template_with_data(html_content: str, form_data: Dict[str, Any], cccd_data: Dict[str, Any] = None) -> str:
    """
    Điền data vào template HTML
    """
    print(f"Filling HTML template with data: {form_data}")
    if cccd_data:
        print(f"CCCD data available: {cccd_data}")
    
    # Parse HTML với BeautifulSoup
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')
    
    replacements_made = 0
    
    # Helper function to get data with fallback to CCCD
    def get_field_value(field_name: str, cccd_field: str = None) -> str:
        # Ưu tiên form_data, fallback sang cccd_data nếu có
        value = form_data.get(field_name, "")
        if not value and cccd_data and cccd_field:
            value = cccd_data.get(cccd_field, "")
        return str(value) if value else ""
    
    # 1. Điền "Kính gửi" - tìm và thay thế chuỗi có dấu ...
    co_quan_tiep_nhan = get_field_value("co_quan_tiep_nhan")
    if co_quan_tiep_nhan:
        # Duyệt mọi span, vì span có thể chứa thẻ con (sup)
        filled_kinh_gui = False
        for span in soup.find_all('span'):
            span_text = span.get_text(strip=False)
            if span_text and "Kính gửi" in span_text:
                # Dựng lại đúng cấu trúc: "Kính gửi" + sup(1) + ": " + cơ quan
                span.clear()
                span.append("Kính gửi")
                sup_tag = soup.new_tag("sup")
                sup_tag.string = "(1)"
                span.append(sup_tag)
                span.append(": ")
                span.append(co_quan_tiep_nhan)
                replacements_made += 1
                filled_kinh_gui = True
                print(f"Filled 'Kính gửi': {co_quan_tiep_nhan}")
                break
        if not filled_kinh_gui:
            print("Không tìm thấy span phù hợp cho 'Kính gửi'")
    
    # 2. Điền "Họ, chữ đệm và tên" - với fallback từ CCCD
    ho_ten_element = soup.find(string=lambda text: text and "1. Họ, chữ đệm và tên:" in text)
    ho_ten = get_field_value("ho_ten", "personName")
    if ho_ten_element and ho_ten:
        new_text = ho_ten_element + " " + ho_ten
        ho_ten_element.replace_with(new_text)
        replacements_made += 1
        print(f"Filled 'Họ tên': {ho_ten}")
    
    # 3. Điền ngày sinh - với fallback từ CCCD
    ngay_sinh_element = soup.find(string=lambda text: text and "2. Ngày, tháng, năm sinh:" in text and "............/............./....................... 3. Giới tính:" in text)
    ngay_sinh = get_field_value("ngay_sinh", "dateOfBirth")
    gioi_tinh = get_field_value("gioi_tinh", "gender")
    
    if ngay_sinh_element and ngay_sinh:
        # Parse ngày sinh - handle both formats
        if "-" in ngay_sinh:  # yyyy-mm-dd (từ form)
            parts = ngay_sinh.split("-")
            day, month, year = parts[2], parts[1], parts[0]
        elif "/" in ngay_sinh:  # dd/mm/yyyy (từ CCCD)
            parts = ngay_sinh.split("/")
            day, month, year = parts[0], parts[1], parts[2] if len(parts) == 3 else ""
        else:
            day, month, year = "", "", ""
        
        new_text = f"2. Ngày, tháng, năm sinh: {day}/{month}/{year} 3. Giới tính: {gioi_tinh}"
        ngay_sinh_element.replace_with(new_text)
        replacements_made += 1
        print(f"Filled 'Ngày sinh': {day}/{month}/{year}, Giới tính: {gioi_tinh}")
    
    # 4. Điền số định danh cá nhân vào các ô - với fallback từ CCCD
    id_boxes = soup.find_all("div", class_="id-box")
    so_dinh_danh = get_field_value("so_dinh_danh", "idCode") or get_field_value("so_cccd", "idCode")
    if so_dinh_danh and len(id_boxes) >= 12:
        so_dinh_danh = str(so_dinh_danh).replace(" ", "").replace("-", "")
        for i, digit in enumerate(so_dinh_danh[:12]):
            if i < len(id_boxes):
                id_boxes[i].string = digit
                replacements_made += 1
        print(f"Filled 'Số định danh': {so_dinh_danh}")
    
    # 5. Điền số điện thoại - tìm với dấu chấm chính xác
    sdt_element = soup.find(string=lambda text: text and "5. Số điện thoại liên hệ: ..............." in text)
    sdt = get_field_value("dien_thoai") or get_field_value("so_dien_thoai")
    if sdt_element and sdt:
        new_text = sdt_element.replace("...............", sdt)
        sdt_element.replace_with(new_text)
        replacements_made += 1
        print(f"Filled 'SĐT': {sdt}")
    
    # 6. Điền email - tìm exact text
    email_element = soup.find(string="6. Email:")
    email = get_field_value("email")
    if email_element and email:
        new_text = f"6. Email: {email}"
        email_element.replace_with(new_text)
        replacements_made += 1
        print(f"Filled 'Email': {email}")
    
    # 6. Điền thông tin chủ hộ - với fallback
    chu_ho_element = soup.find(string=lambda text: text and "7. Họ, chữ đệm và tên chủ hộ:" in text and "8. Mối quan hệ với chủ hộ:" in text)
    if chu_ho_element:
        chu_ho = get_field_value("chu_ho") or get_field_value("ho_ten_chu_ho") or get_field_value("ho_ten", "personName")  # Multiple fallbacks
        quan_he = get_field_value("quan_he_chu_ho") or get_field_value("moi_quan_he_chu_ho", "") or "Chủ hộ"  # Default to "Chủ hộ"
        new_text = f"7. Họ, chữ đệm và tên chủ hộ: {chu_ho} 8. Mối quan hệ với chủ hộ: {quan_he}"
        chu_ho_element.replace_with(new_text)
        replacements_made += 1
        print(f"Filled 'Chủ hộ': {chu_ho}, Quan hệ: {quan_he}")
    
    # 7. Điền số định danh chủ hộ vào các ô thứ 2
    so_dinh_danh_chu_ho = get_field_value("dinh_danh_chu_ho") or get_field_value("so_dinh_danh_chu_ho")
    if so_dinh_danh_chu_ho and len(id_boxes) >= 24:
        dinh_danh_chu_ho = str(so_dinh_danh_chu_ho).replace(" ", "").replace("-", "")
        for i, digit in enumerate(dinh_danh_chu_ho[:12]):
            if i + 12 < len(id_boxes):
                id_boxes[i + 12].string = digit
                replacements_made += 1
        print(f"Filled 'Số định danh chủ hộ': {so_dinh_danh_chu_ho}")
    
    # 10. Điền nội dung đề nghị
    noi_dung_de_nghi = get_field_value("noi_dung_de_nghi")
    if noi_dung_de_nghi:
        noi_dung_elements = soup.find_all(string=lambda text: text and "10. Nội dung đề nghị" in text)
        for noi_dung_element in noi_dung_elements:
            parent = noi_dung_element.parent
            if parent:
                # Thêm nội dung sau element hiện tại
                new_text = f"{noi_dung_element}   {noi_dung_de_nghi}"
                noi_dung_element.replace_with(new_text)
                replacements_made += 1
                print(f"Filled 'Nội dung đề nghị': {noi_dung_de_nghi}")
                break
    

    
    # 9. Điền thông tin thành viên gia đình vào bảng
    members = form_data.get("thanh_vien_gia_dinh") or form_data.get("thanh_vien_ho_gia_dinh") or form_data.get("thanh_vien_cung_thay_doi")
    print(f"Members data: {members}")
    if members:
        table_rows = soup.select("table tbody tr")
        print(f"Found {len(table_rows)} table rows")
        if isinstance(members, list):
            for i, member in enumerate(members[:len(table_rows)]):
                print(f"Processing member {i}: {member}")
                if i < len(table_rows):
                    cells = table_rows[i].find_all("td")
                    if len(cells) >= 6:
                        # STT
                        cells[0].string = str(i + 1)
                        # Họ tên
                        cells[1].string = member.get("ho_ten", "")
                        # Ngày sinh
                        cells[2].string = member.get("ngay_sinh", "")
                        # Giới tính
                        cells[3].string = member.get("gioi_tinh", "")
                        # Số định danh
                        cells[4].string = member.get("so_dinh_danh", "")
                        # Quan hệ với chủ hộ
                        cells[5].string = member.get("moi_quan_he", member.get("quan_he", ""))
                        replacements_made += 6
            print(f"Filled {len(members)} thành viên gia đình")
    
    # 10. Điền thông tin chữ ký (nếu có)
    signature_sections = soup.find_all("div", class_="signature-box")
    
    # Ý kiến chủ hộ
    if len(signature_sections) > 0 and form_data.get("chu_ho_ho_ten"):
        ho_ten_input = signature_sections[0].find(string=lambda text: text and "Họ và tên:" in text)
        if ho_ten_input:
            new_text = ho_ten_input.replace("........................", form_data["chu_ho_ho_ten"])
            ho_ten_input.replace_with(new_text)
            replacements_made += 1
        
        if form_data.get("chu_ho_dinh_danh"):
            dinh_danh_input = signature_sections[0].find(string=lambda text: text and "Số định danh cá nhân:" in text)
            if dinh_danh_input:
                new_text = dinh_danh_input.replace("........................", form_data["chu_ho_dinh_danh"])
                dinh_danh_input.replace_with(new_text)
                replacements_made += 1
    
    # Ý kiến chủ sở hữu chỗ ở
    if len(signature_sections) > 1 and form_data.get("chu_so_huu_ho_ten"):
        ho_ten_input = signature_sections[1].find(string=lambda text: text and "Họ và tên:" in text)
        if ho_ten_input:
            new_text = ho_ten_input.replace("........................", form_data["chu_so_huu_ho_ten"])
            ho_ten_input.replace_with(new_text)
            replacements_made += 1
        
        if form_data.get("chu_so_huu_dinh_danh"):
            dinh_danh_input = signature_sections[1].find(string=lambda text: text and "Số định danh cá nhân:" in text)
            if dinh_danh_input:
                new_text = dinh_danh_input.replace("........................", form_data["chu_so_huu_dinh_danh"])
                dinh_danh_input.replace_with(new_text)
                replacements_made += 1
    
    # Ý kiến cha, mẹ hoặc người giám hộ
    if len(signature_sections) > 2 and form_data.get("giam_ho_ho_ten"):
        ho_ten_input = signature_sections[2].find(string=lambda text: text and "Họ và tên:" in text)
        if ho_ten_input:
            new_text = ho_ten_input.replace("........................", form_data["giam_ho_ho_ten"])
            ho_ten_input.replace_with(new_text)
            replacements_made += 1
        
        if form_data.get("giam_ho_dinh_danh"):
            dinh_danh_input = signature_sections[2].find(string=lambda text: text and "Số định danh cá nhân:" in text)
            if dinh_danh_input:
                new_text = dinh_danh_input.replace("........................", form_data["giam_ho_dinh_danh"])
                dinh_danh_input.replace_with(new_text)
                replacements_made += 1
    
    # Điền số định danh cá nhân vào các ô riêng biệt
    so_dinh_danh = get_field_value("so_dinh_danh", "personalNumber")
    if so_dinh_danh:
        # Tìm div chứa các ô số định danh (4. Số định danh cá nhân)
        id_boxes_container = soup.find('div', class_='id-boxes')
        if id_boxes_container:
            id_boxes = id_boxes_container.find_all('div', class_='id-box')
            so_dinh_danh_str = str(so_dinh_danh).zfill(12)  # Đảm bảo 12 chữ số
            for i, box in enumerate(id_boxes):
                if i < len(so_dinh_danh_str):
                    box.string = so_dinh_danh_str[i]
            replacements_made += 1
            print(f"Filled 'Số định danh cá nhân': {so_dinh_danh}")
    
    # Điền số định danh chủ hộ vào các ô riêng biệt
    so_dinh_danh_chu_ho = get_field_value("so_dinh_danh_chu_ho")
    if so_dinh_danh_chu_ho:
        # Tìm div thứ 2 chứa các ô số định danh (9. Số định danh cá nhân của chủ hộ)
        all_id_containers = soup.find_all('div', class_='id-boxes')
        if len(all_id_containers) > 1:
            chu_ho_container = all_id_containers[1]
            id_boxes = chu_ho_container.find_all('div', class_='id-box')
            so_dinh_danh_str = str(so_dinh_danh_chu_ho).zfill(12)
            for i, box in enumerate(id_boxes):
                if i < len(so_dinh_danh_str):
                    box.string = so_dinh_danh_str[i]
            replacements_made += 1
            print(f"Filled 'Số định danh chủ hộ': {so_dinh_danh_chu_ho}")
    
    # Điền bảng thành viên hộ gia đình
    thanh_vien_list = form_data.get("thanh_vien_ho_gia_dinh", [])
    if thanh_vien_list and isinstance(thanh_vien_list, list):
        table = soup.find('table')
        if table:
            tbody = table.find('tbody')
            if tbody:
                rows = tbody.find_all('tr')
                for i, thanh_vien in enumerate(thanh_vien_list):
                    if i < len(rows):
                        row = rows[i]
                        cells = row.find_all('td')
                        if len(cells) >= 6:
                            # STT
                            cells[0].string = str(i + 1)
                            # Họ tên
                            cells[1].string = thanh_vien.get('ho_ten', '')
                            # Ngày sinh
                            cells[2].string = thanh_vien.get('ngay_sinh', '')
                            # Giới tính
                            cells[3].string = thanh_vien.get('gioi_tinh', '')
                            # Số định danh
                            cells[4].string = thanh_vien.get('so_dinh_danh', '')
                            # Mối quan hệ
                            cells[5].string = thanh_vien.get('moi_quan_he', '')
                replacements_made += len(thanh_vien_list)
                print(f"Filled {len(thanh_vien_list)} thành viên hộ gia đình")
    
    print(f"Total replacements made: {replacements_made}")
    
    # Debug: Save HTML to file for inspection
    with open("debug_filled_template.html", "w", encoding="utf-8") as f:
        f.write(str(soup))
    print("Debug: HTML template saved to debug_filled_template.html")
    
    return str(soup)

def convert_html_to_format(html_content: str, format_type: str) -> bytes:
    """
    Convert HTML to PDF or DOCX with proper formatting
    """
    if format_type == "pdf":
        # Convert HTML to PDF using weasyprint (better CSS support)
        try:
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration
            
            # Create CSS for better Vietnamese font support
            css_content = """
                @page {
                    size: A4;
                    margin: 2cm;
                }
                body {
                    font-family: 'Times New Roman', serif;
                    font-size: 14px;
                    line-height: 1.4;
                }
                .header {
                    text-align: center;
                    margin-bottom: 20px;
                }
                .header h1 {
                    font-size: 16px;
                    font-weight: bold;
                    margin: 5px 0;
                }
                .form-title {
                    font-size: 16px;
                    font-weight: bold;
                    text-align: center;
                    margin: 20px 0;
                }
                .form-field {
                    margin: 10px 0;
                    display: flex;
                    align-items: baseline;
                }
                .form-field label {
                    min-width: 200px;
                    display: inline-block;
                }
                .id-boxes {
                    display: inline-flex;
                    gap: 2px;
                    margin-left: 10px;
                }
                .id-box {
                    width: 22px;
                    height: 22px;
                    border: 1px solid #000;
                    text-align: center;
                    display: inline-block;
                    line-height: 22px;
                    font-size: 12px;
                    font-weight: bold;
                }
                table {
                    width: 100%;
                    border-collapse: collapse;
                    margin: 15px 0;
                    font-size: 12px;
                }
                table, th, td {
                    border: 1px solid #000;
                }
                th, td {
                    padding: 5px;
                    text-align: center;
                    vertical-align: middle;
                    border: 1px solid #000;
                }
                th {
                    background-color: #f0f0f0;
                    font-weight: bold;
                    font-size: 12px;
                }
                td {
                    font-size: 11px;
                }
                .signature-section {
                    display: flex;
                    justify-content: space-between;
                    margin-top: 30px;
                    font-size: 12px;
                }
                .signature-box {
                    width: 23%;
                    border: 1px solid #000;
                    padding: 10px;
                    min-height: 120px;
                    text-align: center;
                }
                .signature-title {
                    font-weight: bold;
                    margin-bottom: 10px;
                }
                .notes {
                    margin-top: 30px;
                    font-size: 12px;
                }
                .notes h3 {
                    font-weight: bold;
                    margin-bottom: 10px;
                }
                .note-item {
                    margin: 10px 0;
                    text-align: justify;
                }
            """
            
            font_config = FontConfiguration()
            css = CSS(string=css_content, font_config=font_config)
            
            # Generate PDF
            html_doc = HTML(string=html_content)
            pdf_bytes = html_doc.write_pdf(stylesheets=[css], font_config=font_config)
            return pdf_bytes
            
        except ImportError as e:
            print(f"weasyprint not available: {e}")
            # Fallback to simple HTML export
            return html_content.encode('utf-8')
        except Exception as e:
            print(f"PDF conversion failed: {e}")
            return html_content.encode('utf-8')
    
    elif format_type == "docx":
        # Convert HTML to DOCX with proper formatting
        try:
            # Skip html2docx and use python-docx directly for better format control
            print("Using python-docx for better format control")
            raise ImportError("Skip html2docx")
            
        except ImportError:
            print("html2docx not available, using python-docx fallback")
            # Fallback to basic docx creation
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from bs4 import BeautifulSoup
                
                doc = Document()
                
                # Set page margins
                section = doc.sections[0]
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(0.8)
                section.right_margin = Inches(0.8)
                
                # Parse HTML and extract structured content
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Add header
                header_elements = soup.find_all(class_='header')
                for header in header_elements:
                    for element in header.find_all(['h1', 'p', 'div']):
                        if element.get_text().strip():
                            p = doc.add_paragraph(element.get_text().strip())
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if element.name == 'h1':
                                for run in p.runs:
                                    run.font.size = Pt(16)
                                    run.bold = True
                
                # Add form title
                form_title = soup.find(class_='form-title')
                if form_title:
                    p = doc.add_paragraph(form_title.get_text().strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(16)
                        run.bold = True
                
                # Add form fields - xử lý riêng từng field để giữ format
                form_fields = soup.find_all(class_='form-field')
                for field in form_fields:
                    field_text = field.get_text().strip()
                    
                    # Xử lý riêng cho số định danh (các ô riêng biệt)
                    if "Số định danh cá nhân" in field_text and "của chủ hộ" not in field_text:
                        # Tách text và số định danh
                        if ":" in field_text:
                            label_part = field_text.split(":")[0] + ":"
                            
                            # Tìm các ô số định danh
                            id_boxes = field.find_all('div', class_='id-box')
                            if id_boxes:
                                # Tạo paragraph với label và table cùng hàng
                                p = doc.add_paragraph()
                                run = p.add_run(label_part + " ")
                                run.font.size = Pt(12)
                                run.font.name = 'Times New Roman'
                                
                                # Tạo các ô cho từng chữ số
                                digits = []
                                for box in id_boxes:
                                    if box.string and box.string.strip():
                                        digits.append(box.string.strip())
                                
                                if digits:
                                    # Tạo table inline cho các ô số
                                    from docx.table import Table
                                    from docx.oxml import OxmlElement
                                    from docx.oxml.ns import qn
                                    
                                    # Tạo table nhỏ inline
                                    table = doc.add_table(rows=1, cols=len(digits))
                                    table.style = 'Table Grid'
                                    
                                    # Set table width nhỏ để không chiếm hết dòng
                                    table.width = Inches(len(digits) * 0.25)  # 0.25 inch per digit
                                    
                                    for i, digit in enumerate(digits):
                                        cell = table.cell(0, i)
                                        cell.text = digit
                                        cell.width = Inches(0.25)
                                        
                                        # Format cell
                                        for paragraph in cell.paragraphs:
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            for run in paragraph.runs:
                                                run.font.size = Pt(10)
                                                run.font.name = 'Times New Roman'
                                                run.bold = True
                                else:
                                    # Tạo ô trống
                                    table = doc.add_table(rows=1, cols=12)
                                    table.style = 'Table Grid'
                                    table.width = Inches(3.0)
                                    
                                    for i in range(12):
                                        cell = table.cell(0, i)
                                        cell.text = ""
                                        cell.width = Inches(0.25)
                                        
                                        # Format cell
                                        for paragraph in cell.paragraphs:
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            for run in paragraph.runs:
                                                run.font.size = Pt(10)
                                                run.font.name = 'Times New Roman'
                            else:
                                # Nếu không tìm thấy id-box, giữ nguyên text
                                p = doc.add_paragraph(field_text)
                                for run in p.runs:
                                    run.font.size = Pt(12)
                                    run.font.name = 'Times New Roman'
                    
                    # Xử lý riêng cho số định danh chủ hộ
                    elif "Số định danh cá nhân của chủ hộ" in field_text:
                        # Tách text và số định danh
                        if ":" in field_text:
                            label_part = field_text.split(":")[0] + ":"
                            
                            # Tìm các ô số định danh
                            id_boxes = field.find_all('div', class_='id-box')
                            if id_boxes:
                                # Tạo paragraph với label và table cùng hàng
                                p = doc.add_paragraph()
                                run = p.add_run(label_part + " ")
                                run.font.size = Pt(12)
                                run.font.name = 'Times New Roman'
                                
                                # Tạo các ô cho từng chữ số
                                digits = []
                                for box in id_boxes:
                                    if box.string and box.string.strip():
                                        digits.append(box.string.strip())
                                
                                if digits:
                                    # Tạo table inline cho các ô số
                                    from docx.table import Table
                                    from docx.oxml import OxmlElement
                                    from docx.oxml.ns import qn
                                    
                                    # Tạo table nhỏ inline
                                    table = doc.add_table(rows=1, cols=len(digits))
                                    table.style = 'Table Grid'
                                    
                                    # Set table width nhỏ để không chiếm hết dòng
                                    table.width = Inches(len(digits) * 0.25)  # 0.25 inch per digit
                                    
                                    for i, digit in enumerate(digits):
                                        cell = table.cell(0, i)
                                        cell.text = digit
                                        cell.width = Inches(0.25)
                                        
                                        # Format cell
                                        for paragraph in cell.paragraphs:
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            for run in paragraph.runs:
                                                run.font.size = Pt(10)
                                                run.font.name = 'Times New Roman'
                                                run.bold = True
                                else:
                                    # Tạo ô trống
                                    table = doc.add_table(rows=1, cols=12)
                                    table.style = 'Table Grid'
                                    table.width = Inches(3.0)
                                    
                                    for i in range(12):
                                        cell = table.cell(0, i)
                                        cell.text = ""
                                        cell.width = Inches(0.25)
                                        
                                        # Format cell
                                        for paragraph in cell.paragraphs:
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            for run in paragraph.runs:
                                                run.font.size = Pt(10)
                                                run.font.name = 'Times New Roman'
                            else:
                                # Nếu không tìm thấy id-box, giữ nguyên text
                                p = doc.add_paragraph(field_text)
                                for run in p.runs:
                                    run.font.size = Pt(12)
                                    run.font.name = 'Times New Roman'
                    
                    else:
                        # Add paragraph with proper formatting for other fields
                        p = doc.add_paragraph(field_text)
                        for run in p.runs:
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'
                
                # Add tables with proper formatting
                tables = soup.find_all('table')
                print(f"Found {len(tables)} tables in HTML")
                for html_table in tables:
                    rows = html_table.find_all('tr')
                    print(f"Table has {len(rows)} rows")
                    if rows:
                        # Calculate max columns, ignoring rowspan/colspan for now
                        max_cols = 6  # Fixed for CT01 table structure
                        
                        # Create table with proper dimensions
                        docx_table = doc.add_table(rows=len(rows), cols=max_cols)
                        docx_table.style = 'Table Grid'
                        
                        # Set table properties for better appearance
                        docx_table.autofit = False
                        
                        # Set table width to fit page
                        from docx.shared import Inches
                        docx_table.width = Inches(6.5)
                        
                        for i, row in enumerate(rows):
                            cells = row.find_all(['th', 'td'])
                            docx_row = docx_table.rows[i]
                            
                            for j, cell in enumerate(cells):
                                if j < len(docx_row.cells):
                                    docx_cell = docx_row.cells[j]
                                    cell_text = cell.get_text().strip()
                                    
                                    # Clear existing content and add new
                                    docx_cell.text = cell_text
                                    
                                    # Format header cells
                                    if cell.name == 'th':
                                        # Make header bold and center-aligned
                                        for paragraph in docx_cell.paragraphs:
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            for run in paragraph.runs:
                                                run.bold = True
                                                run.font.size = Pt(10)
                                    else:
                                        # Format data cells
                                        for paragraph in docx_cell.paragraphs:
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            for run in paragraph.runs:
                                                run.font.size = Pt(10)
                        
                        # Add spacing after table
                        doc.add_paragraph('')
                
                # Add signature section
                signature_section = soup.find(class_='signature-section')
                if signature_section:
                    doc.add_paragraph('')  # Space
                    
                    # Tạo table cho 4 ô ý kiến
                    signature_table = doc.add_table(rows=1, cols=4)
                    signature_table.style = 'Table Grid'
                    signature_table.width = Inches(6.5)
                    
                    signature_boxes = signature_section.find_all(class_='signature-box')
                    for i, box in enumerate(signature_boxes):
                        if i < 4:  # Chỉ xử lý 4 ô
                            cell = signature_table.cell(0, i)
                            cell.text = box.get_text().strip()
                            
                            # Format cho từng paragraph trong cell
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                                    run.font.name = 'Times New Roman'
                                    if "Ý KIẾN" in run.text or "NGƯỜI KÊ KHAI" in run.text:
                                        run.bold = True
                
                # Add notes
                notes = soup.find(class_='notes')
                if notes:
                    doc.add_paragraph('')  # Space
                    for note_item in notes.find_all(class_='note-item'):
                        doc.add_paragraph(note_item.get_text().strip())
                
                output = io.BytesIO()
                doc.save(output)
                output.seek(0)
                return output.getvalue()
            except Exception as e:
                print(f"DOCX conversion failed: {e}")
                return html_content.encode('utf-8')
        except Exception as e:
            print(f"DOCX conversion failed: {e}")
            return html_content.encode('utf-8')
    
    else:
        return html_content.encode('utf-8')